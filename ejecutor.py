import pandas as pd
from datetime import timedelta
from docx import Document
import os
from copy import deepcopy
from configuracion import Configuracion
from funciones import Funciones
from docx.shared import Pt
from docx.shared import Cm
from docx.table import Table


class ejecutor:
    
    
            
    @staticmethod  
    def escribir_profesionales(doc, df_equipo, equipo, tabla_profesional_base):

        from copy import deepcopy
        from docx.table import Table

        profesionales = Configuracion.EQUIPOS[equipo]

        for profesional in profesionales:

            df_prof = df_equipo[df_equipo["PROFESIONAL"] == profesional]

            if df_prof.empty:
                continue

            nombre = Configuracion.NOMBRES_PROFESIONALES.get(
                profesional,
                profesional
            )

            print(f"   👤 {nombre}: {len(df_prof)} actividades")

            # =========================
            # CREAR TABLA DESDE TEMPLATE
            # =========================
            nueva_tabla_xml = deepcopy(tabla_profesional_base)
            doc._body._element.append(nueva_tabla_xml)

            tabla = Table(nueva_tabla_xml, doc)


            # Mantener anchos fijos de la plantilla
            tabla.autofit = False
            tabla.allow_autofit = False


                        
            # =========================
            # NOMBRE PROFESIONAL
            # =========================
            for row in tabla.rows:
                if (
                    len(row.cells) > 1 and
                    "profesional" in row.cells[0].text.lower()
                ):
                    Funciones.escribir_celda(row.cells[1], nombre)
                    break

            # =========================
            # UBICAR CABECERA DE ACTIVIDADES
            # =========================
            header_index = None

            for i, row in enumerate(tabla.rows):

                texto = " ".join(
                    c.text.replace("\n", " ").lower()
                    for c in row.cells
                )

                if "fecha" in texto and "inicio" in texto:
                    header_index = i
                    break

            if header_index is None:
                print(f"❌ No se encontró cabecera para {nombre}")
                continue

            # =========================
            # DEJAR SOLO UNA FILA MODELO
            # (la inmediatamente posterior al header)
            # =========================
            while len(tabla.rows) > header_index + 2:
                tabla._element.remove(
                    tabla.rows[header_index + 2]._element
                )

            fila_modelo = tabla.rows[header_index + 1]
            anchos = [cell.width for cell in fila_modelo.cells]
        
        # =========================
        # OBTENER COLUMNAS REALES
        # =========================
        def obtener_columna(nombre, df_local):

            candidatas = []

            for c in df_local.columns:
                c_norm = (
                    str(c)
                    .upper()
                    .replace("\n", " ")
                    .strip()
                )

                if nombre in c_norm:
                    candidatas.append(c)

            # devolver la primera con datos
            for c in candidatas:
                serie = (
                    df_local[c]
                    .astype(str)
                    .str.strip()
                    .replace(
                        ["", "nan", "None", "---"],
                        pd.NA
                    )
                )

                if serie.notna().any():
                    return c

            return candidatas[0] if candidatas else None

        col_actividad = obtener_columna(
            "ACTIVIDAD", df_prof
        )
        col_detalle = obtener_columna(
            "DETALLE", df_prof
        )
        col_estado = obtener_columna(
            "ESTADO", df_prof
        )
        col_tiempo = obtener_columna(
            "TIEMPO", df_prof
        )

        # =========================
        # COMPLETAR FILA MODELO + AGREGAR RESTO
        # =========================
        for idx, (_, r) in enumerate(df_prof.iterrows()):

            if idx == 0:
                fila = fila_modelo.cells
            else:
                nueva_fila_xml = deepcopy(fila_modelo._tr)
                tabla._tbl.append(nueva_fila_xml)
                fila = tabla.rows[-1].cells
                # Reaplicar ancho original de la plantilla


            fecha = Funciones.formatear_fecha(
                r["FECHA SOLICITUD"]
            )

            actividad = (
                str(r[col_actividad]).strip()
                if col_actividad and pd.notnull(r[col_actividad])
                else "---"
            )

            detalle = (
                str(r[col_detalle]).strip()
                if col_detalle and pd.notnull(r[col_detalle])
                else "---"
            )

            estado = (
                str(r[col_estado]).strip()
                if col_estado and pd.notnull(r[col_estado])
                else "---"
            )

            tiempo = (
                Funciones.normalizar_tiempo(
                    r[col_tiempo]
                )
                if col_tiempo and pd.notnull(r[col_tiempo])
                else "---"
            )

            if actividad.lower() in ("nan", ""):
                actividad = "---"

            if detalle.lower() in ("nan", ""):
                detalle = "---"

            if estado.lower() in ("nan", ""):
                estado = "---"

            if str(tiempo).lower() in ("nan", ""):
                tiempo = "---"

            Funciones.escribir_celda(fila[0], fecha)
            Funciones.escribir_celda(fila[1], fecha)
            Funciones.escribir_celda(fila[2], actividad)
            Funciones.escribir_celda(fila[3], detalle)
            Funciones.escribir_celda(fila[4], estado)
            Funciones.escribir_celda(fila[5], str(tiempo))

        # Espacio entre profesionales
        doc.add_paragraph()

    # =========================    
    @staticmethod    
    def generar_word(df_equipo, equipo, titulo):

        doc = Document(Configuracion.WORD_TEMPLATE)

        Funciones.actualizar_cabecera(doc, equipo)

        ejecutor.escribir_profesionales(doc, df_equipo, equipo)


        nombre = f"Bitácora.docx"
        ruta = os.path.join(Configuracion.OUTPUT_DIR, nombre)

        doc.save(ruta)

        print(f"✅ {ruta}")




    @staticmethod 
    def ejecutar():
        
        os.makedirs(Configuracion.OUTPUT_DIR, exist_ok=True)

        print("🔄 Leyendo Excel...")
        df = Funciones.leer_todos_los_profesionales(
            Configuracion.EXCEL_PATH
        )

        print("🧹 Preparando datos...")
        df = Funciones.preparar_datos(df)

        print("📅 Filtrando por fecha...")
        df = Funciones.filtrar_por_fecha(df)

        print("📄 Generando documento único...")

        doc = Document(Configuracion.WORD_TEMPLATE)

        # =========================
        # TOMAR LAS 3 TABLAS DEL TEMPLATE
        # =========================
        cabecera_base = deepcopy(doc.tables[0]._element)
        tabla_profesional_base = deepcopy(doc.tables[1]._element)
        tabla_final_base = deepcopy(doc.tables[2]._element)

        body = doc._body._element

        # =========================
        # LIMPIAR DOCUMENTO COMPLETO
        # para evitar espacio inicial
        # =========================
        for element in list(body):
            body.remove(element)

        primera_seccion = True

        for equipo in Configuracion.EQUIPOS_ORDEN:

            df_eq = df[df["EQUIPO"] == equipo]

            if df_eq.empty:
                continue

            titulo = Configuracion.MAP_TITULOS[equipo]

            # si MAP_TITULOS tiene tupla/lista usar el primero
            titulo_servicio = (
                titulo[0]
                if isinstance(titulo, (list, tuple))
                else titulo
            )

            print(f"\n📁 Procesando equipo: {titulo_servicio}")

            # =========================
            # SEPARACIÓN ENTRE EQUIPOS
            # =========================
            if not primera_seccion:
                doc.add_paragraph()
                doc.add_page_break()

            primera_seccion = False

            # =========================
            # CABECERA
            # =========================
            nueva_cabecera_xml = deepcopy(cabecera_base)
            body.append(nueva_cabecera_xml)

            tabla_cabecera = Table(nueva_cabecera_xml, doc)

            s1, s2 = Funciones.obtener_semanas(
                Configuracion.FECHA_INICIO,
                Configuracion.FECHA_FIN
            )

            texto_semana = (

                f"Del {Configuracion.FECHA_INICIO} al "
                f"{Configuracion.FECHA_FIN}"
            )

            for row in tabla_cabecera.rows:

                etiqueta = (
                    row.cells[0].text
                    .replace("\n", " ")
                    .strip()
                    .lower()
                )

                # NO tocar contrato
                if etiqueta.startswith("servicio"):
                    Funciones.escribir_celda(
                        row.cells[1],
                        str(titulo_servicio)
                    )

                elif etiqueta.startswith("semana"):
                    Funciones.escribir_celda(
                        row.cells[1],
                        texto_semana
                    )

            # =========================
            # PROFESIONALES
            # =========================
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

            # =========================
            # PROFESIONALES
            # =========================
            ejecutor.escribir_profesionales(
                doc,
                df_eq,
                equipo,
                tabla_profesional_base
            )

            # =========================
            # TABLA FINAL DE HITOS
            # =========================
            nueva_tabla_final_xml = deepcopy(tabla_final_base)
            body.append(nueva_tabla_final_xml)

            # espacio visual después de hitos
            doc.add_paragraph()

        s1, s2 = Funciones.obtener_semanas(
            Configuracion.FECHA_INICIO,
            Configuracion.FECHA_FIN
        )

        nombre = (
            f"Actividades_Semanales.docx"
        )

        ruta = os.path.join(
            Configuracion.OUTPUT_DIR,
            nombre
        )

        doc.save(ruta)

        print(f"\n✅ Documento generado: {ruta}")
        