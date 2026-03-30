import pandas as pd
from datetime import timedelta
from docx import Document
import os
from copy import deepcopy
from configuracion import Configuracion
from funciones import Funciones
from docx.shared import Pt


# =========================
# WORD
# =========================

def escribir_profesionales(doc, df_equipo, equipo, tabla_profesional_base):

    from copy import deepcopy
    from docx.table import Table

    profesionales = Configuracion.EQUIPOS[equipo]

    for profesional in profesionales:

        df_prof = df_equipo[df_equipo["PROFESIONAL"] == profesional]

        if df_prof.empty:
            continue

        nombre = Configuracion.NOMBRES_PROFESIONALES.get(
            profesional,
            profesional
        )

        print(f"   👤 {nombre}: {len(df_prof)} actividades")

        # =========================
        # CREAR TABLA DEL PROFESIONAL
        # =========================
        nueva_tabla_xml = deepcopy(tabla_profesional_base)
        doc._body._element.append(nueva_tabla_xml)

        tabla = Table(nueva_tabla_xml, doc)

        # =========================
        # COMPLETAR NOMBRE PROFESIONAL
        # =========================
        for row in tabla.rows:
            if (
                len(row.cells) > 1 and
                "profesional" in row.cells[0].text.lower()
            ):
                row.cells[1].text = nombre
                break

        # =========================
        # UBICAR HEADER DE ACTIVIDADES
        # =========================
        header_index = None

        for i, row in enumerate(tabla.rows):
            texto = " ".join(
                c.text.replace("\n", " ").lower()
                for c in row.cells
            )

            if "fecha" in texto and "inicio" in texto:
                header_index = i
                break

        if header_index is None:
            print(f"❌ No se encontró cabecera para {nombre}")
            continue

        # =========================
        # LIMPIAR FILAS VACÍAS / EJEMPLO
        # =========================
        while len(tabla.rows) > header_index + 1:
            tabla._element.remove(tabla.rows[header_index + 1]._element)

        # =========================
        # NORMALIZAR COLUMNAS
        # =========================
        columnas = {
            c.upper().replace("\n", " ").strip(): c
            for c in df_prof.columns
        }

        def obtener_columna(nombre):
            for k, v in columnas.items():
                if nombre in k:
                    return v
            return None

        col_estado = Funciones.obtener_columna("ESTADO", df_prof)
        col_actividad = Funciones.obtener_columna("ACTIVIDAD", df_prof)
        col_detalle = Funciones.obtener_columna("DETALLE", df_prof)
        col_tiempo = Funciones.obtener_columna("TIEMPO", df_prof)

        
        # =========================
        # INSERTAR ACTIVIDADES
        # =========================
        for _, r in df_prof.iterrows():

            fila = tabla.add_row().cells

            fi = r["FECHA SOLICITUD"]

            actividad = (
                str(r[col_actividad]).strip()
                if col_actividad and pd.notnull(r[col_actividad])
                else "---"
            )

            detalle = (
                str(r[col_detalle]).strip()
                if col_detalle and pd.notnull(r[col_detalle])
                else "---"
            )

            estado = (
                str(r[col_estado]).strip()
                if col_estado and pd.notnull(r[col_estado])
                else "---"
            )

            tiempo = Funciones.normalizar_tiempo(
                r[col_tiempo] if col_tiempo else None
            )

            if actividad.lower() == "nan":
                actividad = "---"

            if detalle.lower() == "nan":
                detalle = "---"

            if estado.lower() == "nan":
                estado = "---"

            if str(tiempo).lower() == "nan":
                tiempo = "---"

            fila[0].text = Funciones.formatear_fecha(fi)
            fila[1].text = Funciones.formatear_fecha(fi)
            fila[2].text = actividad
            fila[3].text = detalle
            fila[4].text = estado
            fila[5].text = str(tiempo)

        # espacio entre profesionales
        doc.add_paragraph()


# =========================        
def generar_word(df_equipo, equipo, titulo):

    doc = Document(Configuracion.WORD_TEMPLATE)

    Funciones.actualizar_cabecera(doc, equipo)

    escribir_profesionales(doc, df_equipo, equipo)

    s1, s2 = Funciones.obtener_semanas(
        Configuracion.FECHA_INICIO,
        Configuracion.FECHA_FIN
    )

    nombre = f"Actividades_Semanales_{titulo}_semana_{s1}-{s2}.docx"
    ruta = os.path.join(Configuracion.OUTPUT_DIR, nombre)

    doc.save(ruta)

    print(f"✅ {ruta}")


# =========================
# MAIN
# =========================

def main():

    from copy import deepcopy
    from docx import Document
    from docx.table import Table

    os.makedirs(Configuracion.OUTPUT_DIR, exist_ok=True)

    print("🔄 Leyendo Excel...")
    df = Funciones.leer_todos_los_profesionales(
        Configuracion.EXCEL_PATH
    )

    print("🧹 Preparando datos...")
    df = Funciones.preparar_datos(df)

    print("📅 Filtrando por fecha...")
    df = Funciones.filtrar_por_fecha(df)

    print("📄 Generando documento único...")

    doc = Document(Configuracion.WORD_TEMPLATE)

    # Guardar las 3 tablas base de la plantilla
    cabecera_base = deepcopy(doc.tables[0]._element)
    tabla_profesional_base = deepcopy(doc.tables[1]._element)
    tabla_final_base = deepcopy(doc.tables[2]._element)

    body = doc._body._element

    # Eliminar TODO el contenido original del body del documento
    # Esto evita que queden párrafos vacíos al inicio.
    for element in list(body):
        body.remove(element)

    primera_seccion = True

    for equipo in Configuracion.EQUIPOS_ORDEN:

        df_eq = df[df["EQUIPO"] == equipo]

        if df_eq.empty:
            print(f"⚠️ Sin datos para {equipo}")
            continue

        titulo = Configuracion.MAP_TITULOS[equipo]

        # Si MAP_TITULOS trae tupla/lista, usar el primer valor
        titulo_servicio = (
            titulo[0]
            if isinstance(titulo, (tuple, list))
            else titulo
        )

        print(f"📁 Procesando: {titulo_servicio}")

        # Desde el segundo equipo en adelante:
        # dejar espacio, luego salto de página
        if not primera_seccion:
            doc.add_paragraph()
            doc.add_page_break()

        primera_seccion = False

        # =========================
        # TABLA CABECERA
        # =========================
        nueva_cabecera = deepcopy(cabecera_base)
        body.append(nueva_cabecera)

        tabla_cabecera = Table(nueva_cabecera, doc)

        s1, s2 = Funciones.obtener_semanas(
            Configuracion.FECHA_INICIO,
            Configuracion.FECHA_FIN
        )

        texto_semana = (
            f"Parte de la {s1}° y {s2}° semana de "
            f"{Configuracion.MES}\n"
            f"({Configuracion.FECHA_INICIO} al "
            f"{Configuracion.FECHA_FIN})"
        )

        for row in tabla_cabecera.rows:

            etiqueta = (
                row.cells[0].text
                .replace("\n", " ")
                .strip()
                .lower()
            )

            # SOLO modificar Servicio
            if etiqueta.startswith("servicio"):
                row.cells[1].text = str(titulo_servicio)

            # SOLO modificar Semana
            elif etiqueta.startswith("semana"):
                row.cells[1].text = texto_semana

            # Contrato y Administrador quedan intactos desde template

        # =========================
        # TABLAS POR PROFESIONAL
        # =========================
        escribir_profesionales(
            doc,
            df_eq,
            equipo,
            tabla_profesional_base
        )

        # =========================
        # TABLA FINAL HITOS
        # =========================
        nueva_tabla_final = deepcopy(tabla_final_base)
        body.append(nueva_tabla_final)

        # Espacio visual DESPUÉS de la tabla de hitos
        doc.add_paragraph()

    s1, s2 = Funciones.obtener_semanas(
        Configuracion.FECHA_INICIO,
        Configuracion.FECHA_FIN
    )

    nombre = (
        f"Actividades_Semanales_Todos_semana_{s1}-{s2}.docx"
    )

    ruta = os.path.join(
        Configuracion.OUTPUT_DIR,
        nombre
    )

    doc.save(ruta)

    print(f"✅ Documento generado: {ruta}")
      
if __name__ == "__main__":
    main()